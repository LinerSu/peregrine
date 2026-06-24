"""Extract plain text from an uploaded document (PDF / DOCX / .txt / .md).

Shared by CV intake and job-posting ingestion so a user can hand us a saved page
or résumé as a file instead of a URL we'd have to scrape.
"""
from __future__ import annotations

import io


def extract_text(filename: str, raw: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if name.endswith(".docx"):
        # Word résumés are common; pull paragraph + table cell text.
        from docx import Document

        doc = Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    return raw.decode("utf-8", errors="ignore")
